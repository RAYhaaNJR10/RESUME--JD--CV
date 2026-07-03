import os

from docx import Document
from docx.shared import Inches


def generate_docx(
    cv_json,
    output_path,
    template_path=None
):
    if not template_path:
        template_path = "uploads/templates/active_template.docx"
    if template_path and os.path.exists(template_path):
        document = Document(template_path)
    else:
        document = Document()

    # ==========================================
    # Candidate Details
    # ==========================================

    document.add_heading(
        "Candidate Details",
        level=1
    )

    document.add_paragraph(
        f"Candidates Full Name: "
        f"{cv_json.get('candidate_name', '')}"
    )

    current_role = cv_json.get(
        "current_role_title",
        ""
    )

    if not current_role:

        experiences = cv_json.get(
            "professional_experience",
            []
        )

        if experiences:

            current_role = experiences[0].get(
                "role",
                ""
            )

    document.add_paragraph(
        f"Candidates Current Role Title: "
        f"{current_role}"
    )

    # ==========================================
    # About Candidate
    # ==========================================

    about = cv_json.get(
        "about_candidate",
        ""
    )

    if about:

        document.add_heading(
            "About the Candidate",
            level=1
        )

        document.add_paragraph(
            about
        )

    # ==========================================
    # Strengths
    # ==========================================

    strengths = cv_json.get(
        "strengths",
        []
    )

    if strengths:

        document.add_heading(
            "Strengths",
            level=1
        )

        table = document.add_table(
            rows=1,
            cols=3
        )

        table.style = "Table Grid"

        header = table.rows[0].cells

        header[0].text = "Strength"
        header[1].text = "Score (1-5)"
        header[2].text = "Assessment"

        # Set header cell widths
        header[0].width = Inches(2.0)
        header[1].width = Inches(1.0)
        header[2].width = Inches(3.5)

        for strength in strengths:

            row = table.add_row().cells

            # Set row cell widths
            row[0].width = Inches(2.0)
            row[1].width = Inches(1.0)
            row[2].width = Inches(3.5)

            row[0].text = str(
                strength.get(
                    "strength",
                    ""
                )
            )

            row[1].text = str(
                strength.get(
                    "score",
                    ""
                )
            )

            row[2].text = str(
                strength.get(
                    "assessment",
                    ""
                )
            )

    # ==========================================
    # Experience
    # ==========================================

    experiences = cv_json.get(
        "professional_experience",
        []
    )

    for idx, exp in enumerate(
        experiences,
        start=1
    ):

        document.add_heading(
            f"Experience Part {idx}",
            level=1
        )

        document.add_paragraph(
            f"Name of company/client: "
            f"{exp.get('company', '')}"
        )

        document.add_paragraph(
            f"Project Name: "
            f"{exp.get('project', '')}"
        )

        document.add_paragraph(
            f"Role Title performed: "
            f"{exp.get('role', '')}"
        )

        document.add_paragraph(
            "Role Summary:"
        )

        document.add_paragraph(
            exp.get(
                "role_summary",
                ""
            )
        )

        achievements = exp.get(
            "key_achievements",
            []
        )

        if achievements:

            document.add_paragraph(
                "Key Achievements:"
            )

            for achievement in achievements:

                document.add_paragraph(
                    achievement,
                    style="List Bullet"
                )

    # ==========================================
    # Technical Skills
    # ==========================================

    technical_skills = cv_json.get(
        "technical_skills",
        []
    )

    if technical_skills:

        document.add_heading(
            "Technical Skills",
            level=1
        )

        table = document.add_table(
            rows=1,
            cols=2
        )

        table.style = "Table Grid"

        header = table.rows[0].cells

        header[0].text = "Technology Area"
        header[1].text = "Skills"

        for item in technical_skills:

            row = table.add_row().cells

            row[0].text = str(
                item.get(
                    "technology_area",
                    ""
                )
            )

            row[1].text = ", ".join(
                item.get(
                    "skills",
                    []
                )
            )

    # ==========================================
    # Education
    # ==========================================

    education = cv_json.get(
        "education",
        []
    )

    if education:

        document.add_heading(
            "Education",
            level=1
        )

        for edu in education:

            degree = edu.get(
                "degree",
                ""
            )

            institution = edu.get(
                "institution",
                ""
            )

            start_year = str(
                edu.get(
                    "start_year",
                    ""
                )
            )

            end_year = str(
                edu.get(
                    "end_year",
                    ""
                )
            )

            document.add_paragraph(
                degree
            )

            if start_year or end_year:

                document.add_paragraph(
                    f"{institution} | "
                    f"{start_year} - {end_year}"
                )

            else:

                document.add_paragraph(
                    institution
                )

    # ==========================================
    # Save
    # ==========================================

    os.makedirs(
        os.path.dirname(
            output_path
        ),
        exist_ok=True
    )

    document.save(
        output_path
    )

    return output_path